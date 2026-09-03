from __future__ import annotations

import re
from collections.abc import Iterable
from datetime import date
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from models import Change, Lesson
from utils import clean_teacher, clean_text, normalize_key, parse_date_from_text


DAY_ALIASES = {
    "пн": "пн",
    "понедельник": "пн",
    "вт": "вт",
    "вторник": "вт",
    "ср": "ср",
    "среда": "ср",
    "ср. ": "ср",
    "чт": "чт",
    "четверг": "чт",
    "пт": "пт",
    "пятница": "пт",
    "сб": "сб",
    "суббота": "сб",
    "вс": "вс",
    "воскресенье": "вс",
}


def _day_key(value: object) -> str | None:
    text = clean_text(value).casefold().replace(".", "")
    return DAY_ALIASES.get(text)


def parse_xlsx(path: str | Path) -> tuple[list[str], dict[str, dict[str, dict[int, Lesson]]]]:
    """Parse the college's wide Excel layout.

    The source has one group per two columns. Every pair in a weekday block
    occupies two rows: subject/room and teacher. The first column contains a
    merged weekday label.
    """
    workbook = load_workbook(path, data_only=True, read_only=False)
    worksheet = next(
        (sheet for sheet in workbook.worksheets if "распис" in sheet.title.casefold()),
        workbook.worksheets[0],
    )

    group_columns: list[tuple[int, str, str]] = []
    groups: list[str] = []
    for column in range(2, worksheet.max_column + 1, 2):
        group = clean_text(worksheet.cell(4, column).value)
        if not group:
            continue
        group_key = normalize_key(group)
        group_columns.append((column, group, group_key))
        groups.append(group)

    day_starts: list[tuple[int, str]] = []
    for row in range(1, worksheet.max_row + 1):
        day = _day_key(worksheet.cell(row, 1).value)
        if day is not None:
            day_starts.append((row, day))

    lessons: dict[str, dict[str, dict[int, Lesson]]] = {}
    for index, (start_row, day) in enumerate(day_starts):
        stop_row = day_starts[index + 1][0] if index + 1 < len(day_starts) else worksheet.max_row + 1
        slot_count = max(0, (stop_row - start_row) // 2)
        day_data = lessons.setdefault(day, {})

        for slot in range(slot_count):
            subject_row = start_row + slot * 2
            teacher_row = subject_row + 1
            pair_number = slot + 1
            for column, _group, group_key in group_columns:
                subject = clean_text(worksheet.cell(subject_row, column).value)
                room = clean_text(worksheet.cell(subject_row, column + 1).value)
                teacher = clean_teacher(worksheet.cell(teacher_row, column).value)
                if not subject and not room and not teacher:
                    continue
                day_data.setdefault(group_key, {})[pair_number] = Lesson(
                    subject=subject or "Занятие",
                    room=room,
                    teacher=teacher,
                )

    return groups, lessons


def _header_index(headers: list[str], *needles: str, default: int) -> int:
    for index, header in enumerate(headers):
        lowered = header.casefold()
        if any(needle in lowered for needle in needles):
            return index
    return default


def _parse_pair(value: str) -> int | None:
    match = re.search(r"\d+", value)
    return int(match.group()) if match else None


def _split_replacement(value: str) -> tuple[str, str]:
    # The source uses an en dash between the old and new entries.
    parts = re.split(r"\s*[–—]\s*", value, maxsplit=1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    parts = re.split(r"\s+-\s+", value, maxsplit=1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    return "", value.strip()


def _replacement_lesson(new_description: str, room: str, old_description: str) -> Change:
    if re.search(r"\b(снятие|отмена|отменено|отменить)\b", new_description.casefold()):
        return Change(lesson=None, cancelled=True, old_description=old_description)

    teacher = ""
    subject = new_description
    match = re.match(r"^(.*?)\s*\(([^()]*)\)\s*$", new_description)
    if match:
        subject = match.group(1).strip()
        teacher = clean_teacher(match.group(2))

    return Change(
        lesson=Lesson(subject=subject or "Занятие", room=room, teacher=teacher),
        old_description=old_description,
    )


def parse_docx_changes(path: str | Path, default_date: date | None = None) -> dict[date, dict[str, dict[int, Change]]]:
    """Parse a DOCX change notice into date -> group -> pair -> change."""
    document = Document(path)
    paragraphs = " ".join(clean_text(paragraph.text) for paragraph in document.paragraphs)
    change_date = parse_date_from_text(paragraphs) or parse_date_from_text(Path(path).name) or default_date
    if change_date is None:
        return {}

    result: dict[date, dict[str, dict[int, Change]]] = {}
    for table in document.tables:
        if not table.rows:
            continue
        headers = [clean_text(cell.text) for cell in table.rows[0].cells]
        pair_index = _header_index(headers, "пар", default=0)
        group_index = _header_index(headers, "груп", default=1)
        subject_index = _header_index(headers, "предмет", "препод", default=2)
        room_index = _header_index(headers, "кабин", "аудитор", default=3)

        for row in table.rows[1:]:
            cells = [clean_text(cell.text) for cell in row.cells]
            if len(cells) <= max(pair_index, group_index, subject_index):
                continue
            pair = _parse_pair(cells[pair_index])
            group = cells[group_index]
            description = cells[subject_index]
            room = cells[room_index] if room_index < len(cells) else ""
            if pair is None or not group or not description:
                continue

            old_description, new_description = _split_replacement(description)
            change = _replacement_lesson(new_description, room, old_description)
            result.setdefault(change_date, {}).setdefault(normalize_key(group), {})[pair] = change

    return result


def merge_changes(
    destination: dict[date, dict[str, dict[int, Change]]],
    source: dict[date, dict[str, dict[int, Change]]],
) -> None:
    """Merge notices in file order; later files override earlier duplicates."""
    for change_date, groups in source.items():
        for group_key, pairs in groups.items():
            destination.setdefault(change_date, {}).setdefault(group_key, {}).update(pairs)


def parse_change_files(paths: Iterable[str | Path], default_date: date | None = None) -> dict[date, dict[str, dict[int, Change]]]:
    changes: dict[date, dict[str, dict[int, Change]]] = {}
    for path in paths:
        merge_changes(changes, parse_docx_changes(path, default_date=default_date))
    return changes
