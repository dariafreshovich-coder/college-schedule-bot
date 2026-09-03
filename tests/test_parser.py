from __future__ import annotations

from datetime import date

from docx import Document
from openpyxl import Workbook

from models import ScheduleData
from parser import parse_docx_changes, parse_xlsx


def test_excel_and_change_notice_are_merged(tmp_path):
    xlsx_path = tmp_path / "schedule.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Расписание"
    worksheet["B4"] = "ТЕСТ-1"
    worksheet["A5"] = "пт"
    worksheet["B5"] = "Математика"
    worksheet["C5"] = "12"
    worksheet["B6"] = "Иванов И.И."
    workbook.save(xlsx_path)

    docx_path = tmp_path / "замена на 04.09.2026.docx"
    document = Document()
    document.add_paragraph("Изменение в расписании на 04.09.2026 г.")
    table = document.add_table(rows=2, cols=4)
    for cell, value in zip(table.rows[0].cells, ("Пары", "Группа", "Предмет, преподаватель", "Кабинет")):
        cell.text = value
    for cell, value in zip(
        table.rows[1].cells,
        ("1", "ТЕСТ-1", "Математика (Иванов И.И.) – Физика (Петров П.П.)", "27"),
    ):
        cell.text = value
    document.save(docx_path)

    groups, lessons = parse_xlsx(xlsx_path)
    changes = parse_docx_changes(docx_path)
    data = ScheduleData(groups, lessons, changes)
    result = data.get_lessons("тест-1", date(2026, 9, 4))

    assert groups == ["ТЕСТ-1"]
    assert result[1].subject == "Физика"
    assert result[1].teacher == "Петров П.П."
    assert result[1].room == "27"
    assert result[1].changed is True


def test_cancelled_pair(tmp_path):
    path = tmp_path / "change.docx"
    document = Document()
    document.add_paragraph("Изменение на 04.09.2026")
    table = document.add_table(rows=2, cols=4)
    for cell, value in zip(table.rows[0].cells, ("Пары", "Группа", "Предмет", "Кабинет")):
        cell.text = value
    for cell, value in zip(table.rows[1].cells, ("2", "ТЕСТ-1", "Физика (Петров П.П.) – снятие", "")):
        cell.text = value
    document.save(path)

    changes = parse_docx_changes(path)
    assert changes[date(2026, 9, 4)]["тест-1"][2].cancelled is True
